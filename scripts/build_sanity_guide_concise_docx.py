from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "Sanity后台管理操作指南.md"
OUTPUT_DOCX = ROOT / "Sanity后台管理操作指南_精简版.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(94, 103, 115)
LIGHT_BLUE = "E8EEF5"
NOTE_FILL = "F4F6F9"
TABLE_BORDER = "D7DEE8"


@dataclass
class Subsection:
    title: str
    lines: list[str] = field(default_factory=list)


@dataclass
class Section:
    title: str
    lines: list[str] = field(default_factory=list)
    subsections: list[Subsection] = field(default_factory=list)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def add_inline_runs(paragraph, text: str, base_size: float = 10.5, bold_default: bool = False) -> None:
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    text = text.replace("**", "").replace("`", "").replace("\\|", "|").strip()
    run = paragraph.add_run(text)
    set_run_font(run, size=base_size, bold=bold_default)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.22
    add_inline_runs(p, text, base_size=9.7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), NOTE_FILL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "C9D6E2")
    p_bdr.append(left)
    p_pr.append(p_bdr)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("Sanity 后台管理操作指南 · 精简版")
    set_run_font(r, size=8.5, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("天悦宝贝 · 后台内容维护速查")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SANITY")
    set_run_font(r, size=15, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("后台管理操作指南")
    set_run_font(r, size=27, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("精简操作版")
    set_run_font(r, size=16, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("常用路径 · 必填字段 · 图片尺寸 · 发布检查")
    set_run_font(r, size=11, color=MUTED)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("最后更新：2026-07-08")
    set_run_font(r, size=10.5, bold=True, color=MUTED)
    doc.add_page_break()


def parse_sections(lines: list[str]) -> list[Section]:
    sections: list[Section] = []
    current: Section | None = None
    current_sub: Subsection | None = None
    in_toc = False
    for raw in lines:
        line = raw.rstrip()
        h2 = re.match(r"^##\s+(.+)$", line)
        h3 = re.match(r"^###\s+(.+)$", line)
        h4 = re.match(r"^####\s+(.+)$", line)
        if h2:
            title = h2.group(1).strip()
            if title == "目录":
                in_toc = True
                continue
            in_toc = False
            current = Section(title=title)
            sections.append(current)
            current_sub = None
            continue
        if in_toc:
            continue
        if h3 or h4:
            if current is None:
                continue
            title = (h3 or h4).group(1).strip()
            current_sub = Subsection(title=title)
            current.subsections.append(current_sub)
            continue
        if current_sub is not None:
            current_sub.lines.append(line)
        elif current is not None:
            current.lines.append(line)
    return sections


def parse_tables(lines: list[str]) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            if len(rows) > 1:
                tables.append(rows)
            continue
        i += 1
    return tables


def metadata_from_lines(lines: list[str]) -> str:
    notes = []
    for line in lines:
        s = line.strip()
        if s.startswith(">"):
            notes.append(s.lstrip(">").strip().replace("**", "").replace("`", ""))
    return "；".join(notes[:3])


def table_summary(lines: list[str]) -> tuple[str, str]:
    tables = parse_tables(lines)
    fields: list[str] = []
    warnings: list[str] = []
    for table in tables[:2]:
        header = table[0]
        rows = table[1:]
        for row in rows:
            if not row:
                continue
            field_name = re.sub(r"\*\*|`", "", row[0]).strip()
            if field_name and field_name not in fields:
                fields.append(field_name)
            row_text = " / ".join(row)
            if re.search(r"必填|推荐|尺寸|注意|不可|需|默认|可选", row_text):
                clean = re.sub(r"\*\*|`", "", row_text)
                warnings.append(clean)
    field_text = "、".join(fields[:8])
    if len(fields) > 8:
        field_text += " 等"
    warning_text = "；".join(warnings[:3])
    return field_text or "按后台表单填写", warning_text or "修改后记得发布并检查前台"


def first_steps(lines: list[str], limit: int = 4) -> list[str]:
    items = []
    for line in lines:
        s = line.strip()
        m = re.match(r"^(?:\d+\.|-)\s+(.+)$", s)
        if m:
            item = re.sub(r"\*\*|`", "", m.group(1)).strip()
            if item and item not in items:
                items.append(item)
        if len(items) >= limit:
            break
    return items


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    for c, text in enumerate(headers):
        cell = table.rows[0].cells[c]
        shade_cell(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, text, base_size=9.3, bold_default=True)
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, text, base_size=9.0, bold_default=(c == 0))
    if widths:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width * 1440)))
            grid.append(col)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_toc(doc: Document, sections: list[Section]) -> None:
    doc.add_heading("使用前先看", level=1)
    add_note(doc, "这是一份精简操作版，适合日常内容维护快速查找。完整字段仍以后台表单和 Markdown 原文为准。")
    rows = []
    for section in sections:
        if section.title == "目录":
            continue
        subs = "、".join(s.title for s in section.subsections[:5])
        if len(section.subsections) > 5:
            subs += " 等"
        rows.append([section.title, subs or "页面基础配置"])
    add_simple_table(doc, ["章节", "主要内容"], rows, widths=[2.0, 4.5])


def add_quick_start(doc: Document, sections: list[Section]) -> None:
    doc.add_heading("一、通用操作", level=1)
    source = next((s for s in sections if "通用说明" in s.title), None)
    add_note(doc, "后台地址：https://tiancibaobei.sanity.studio；修改内容后必须点击 Publish，前台通常 1 分钟内更新。")
    if source:
        for sub in source.subsections:
            title = sub.title
            if any(key in title for key in ("登录", "界面", "发布", "图片上传")):
                doc.add_heading(title, level=2)
                steps = first_steps(sub.lines, limit=5)
                if steps:
                    for step in steps:
                        p = doc.add_paragraph(style="List Number" if "登录" in title or "发布" in title else "List Bullet")
                        add_inline_runs(p, step, base_size=10.2)
                else:
                    field_text, warning_text = table_summary(sub.lines)
                    add_note(doc, f"{field_text}；{warning_text}")
    add_simple_table(
        doc,
        ["常用组件", "填写重点"],
        [
            ["SEO 设置", "标题 50-60 字符；描述 120-160 字符；分享图推荐 1200×630 px。"],
            ["Banner", "桌面图必填；常用尺寸 2400×800 px；移动图建议 900×600 px。"],
            ["CTA 咨询区", "确认标题、说明、按钮文字和按钮链接，链接常用 /start-your-journey。"],
            ["图片上传", "优先 WebP/JPG；文件小于建议大小；命名用英文，避免特殊字符。"],
        ],
        widths=[1.5, 5.0],
    )


def add_section_digest(doc: Document, section: Section) -> None:
    doc.add_heading(section.title, level=1)
    meta = metadata_from_lines(section.lines)
    if meta:
        add_note(doc, meta)
    rows = []
    subsections = section.subsections or [Subsection("页面基础配置", section.lines)]
    for sub in subsections:
        fields, warning = table_summary(sub.lines)
        steps = first_steps(sub.lines, limit=2)
        action = "；".join(steps) if steps else fields
        rows.append([sub.title, action[:95], warning[:110]])
    if not rows:
        rows.append(["页面配置", "按后台表单填写", "修改后发布并检查前台"])
    add_simple_table(doc, ["模块", "主要维护内容", "注意事项"], rows, widths=[1.55, 2.55, 2.4])


def add_image_appendix(doc: Document, sections: list[Section]) -> None:
    appendix = next((s for s in sections if "图片尺寸" in s.title), None)
    if not appendix:
        return
    doc.add_heading("附录：图片尺寸速查", level=1)
    tables = parse_tables(appendix.lines)
    if tables:
        table = tables[0]
        add_simple_table(doc, table[0], table[1:], widths=[2.0, 1.65, 1.45, 1.4])
    add_note(doc, "遇到问题：先检查是否已 Publish、图片大小和格式是否符合要求；字段不确定时，可在完整 Markdown 中搜索字段名称。")


def build() -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    sections = parse_sections(lines)
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_toc(doc, sections)
    add_quick_start(doc, sections)
    for section in sections:
        if "通用说明" in section.title or "图片尺寸" in section.title:
            continue
        add_section_digest(doc, section)
    add_image_appendix(doc, sections)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
