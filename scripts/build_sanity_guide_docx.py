from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "Sanity后台管理操作指南.md"
REFERENCE_DOCX = Path("D:/photo/康贝儿/Sanity后台管理操作指南2.0_可见目录版.docx")
OUTPUT_DOCX = ROOT / "Sanity后台管理操作指南_直观目录版.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 39, 55)
MUTED = RGBColor(92, 103, 115)
LIGHT_BLUE = "E8EEF5"
NOTE_FILL = "F4F6F9"
NOTE_BORDER = "C9D6E2"
TABLE_BORDER = "D7DEE8"


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_border_bottom(paragraph, color: str = "D6DFEA", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.18
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Note Box" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Note Box", 1)
    note = doc.styles["Note Box"]
    note.font.name = FONT_LATIN
    note._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor(55, 65, 81)
    note.paragraph_format.left_indent = Inches(0.12)
    note.paragraph_format.right_indent = Inches(0.08)
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.25


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("Sanity 后台管理操作指南")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("天悦宝贝 · 内容维护手册")
    set_run_font(r, size=8.5, color=MUTED)


def add_inline_runs(paragraph, text: str, base_size: float = 10.5, bold_default: bool = False) -> None:
    text = text.replace("\\|", "|").strip()
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, bold=bold_default)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, bold=bold_default)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Note Box")
    if "\n" in text:
        parts = text.split("\n")
        for idx, part in enumerate(parts):
            add_inline_runs(p, part, base_size=10)
            if idx < len(parts) - 1:
                p.add_run().add_break()
    else:
        add_inline_runs(p, text, base_size=10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), NOTE_FILL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), NOTE_BORDER)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SANITY")
    set_run_font(r, size=15, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("后台管理操作指南")
    set_run_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("直观目录版")
    set_run_font(r, size=16, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("网站内容维护 · 页面配置 · 图片规范 · 咨询线索")
    set_run_font(r, size=11, color=MUTED)

    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("最后更新：2026-07-07")
    set_run_font(r, size=10.5, bold=True, color=MUTED)

    doc.add_page_break()


def parse_heading(line: str) -> tuple[int, str] | None:
    m = re.match(r"^(#{1,4})\s+(.+)$", line)
    if not m:
        return None
    level = len(m.group(1))
    title = re.sub(r"\s*\{#.+?\}\s*$", "", m.group(2)).strip()
    return level, title


def get_headings(lines: list[str]) -> list[tuple[int, str]]:
    result: list[tuple[int, str]] = []
    in_md_toc = False
    for line in lines:
        h = parse_heading(line)
        if h and h[1] == "目录":
            in_md_toc = True
            continue
        if in_md_toc and line.strip() == "---":
            in_md_toc = False
            continue
        if in_md_toc:
            continue
        if h and h[0] > 1:
            result.append(h)
    return result


def add_visible_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("使用前先看")
    set_run_font(r, size=16, bold=True, color=BLUE)
    set_paragraph_border_bottom(p)
    add_note(doc, "这份手册适合日常内容维护人员。先从目录找到对应模块，再按字段说明填写；涉及图片时，请优先查看文末“图片尺寸速查表”。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("目录")
    set_run_font(r, size=16, bold=True, color=BLUE)
    set_paragraph_border_bottom(p)
    add_note(doc, "以下目录为可见目录，打开文档即可查看，无需右键更新。")

    for level, title in headings:
        if level > 3:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.0 if level == 2 else 0.22)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(2)
        text = title
        r = p.add_run(text)
        set_run_font(r, size=10.5 if level == 2 else 9.5, bold=(level == 2),
                     color=INK if level == 2 else MUTED)
    doc.add_page_break()


def is_md_table_start(lines: list[str], i: int) -> bool:
    return (
        i + 1 < len(lines)
        and lines[i].strip().startswith("|")
        and lines[i + 1].strip().startswith("|")
        and set(lines[i + 1].strip().replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set()
    )


def parse_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [9360 // cols for _ in range(cols)]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(6.5 / cols)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                shade_cell(cell, LIGHT_BLUE)
            text = rows[row_idx][col_idx] if col_idx < len(rows[row_idx]) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            add_inline_runs(p, text, base_size=9.2, bold_default=(row_idx == 0))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown_body(doc: Document, lines: list[str]) -> None:
    in_md_toc = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        heading = parse_heading(stripped)
        if heading and heading[1] == "目录":
            in_md_toc = True
            i += 1
            continue
        if in_md_toc:
            if stripped == "---":
                in_md_toc = False
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if heading:
            level, title = heading
            if level == 1:
                i += 1
                continue
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_inline_runs(p, title, base_size=16 if style == "Heading 1" else 13 if style == "Heading 2" else 12,
                            bold_default=True)
            i += 1
            continue

        if stripped.startswith(">"):
            note_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                note_lines.append(lines[i].strip().lstrip(">").strip().rstrip())
                i += 1
            add_note(doc, "\n".join(note_lines))
            continue

        if is_md_table_start(lines, i):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if numbered or bullet:
            style = "List Number" if numbered else "List Bullet"
            text = (numbered or bullet).group(1)
            p = doc.add_paragraph(style=style)
            add_inline_runs(p, text)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1


def build() -> None:
    doc = Document(str(REFERENCE_DOCX)) if REFERENCE_DOCX.exists() else Document()
    clear_document_body(doc)
    configure_styles(doc)
    setup_page(doc)

    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    headings = get_headings(lines)

    add_cover(doc)
    add_visible_toc(doc, headings)
    add_markdown_body(doc, lines)

    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
