from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "tmp_home_compare"
OUT_DOCX = ROOT / "Sanity后台首页截图对照样稿.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(36, 101, 168)
DARK = RGBColor(28, 35, 45)
MUTED = RGBColor(96, 107, 122)
LIGHT_BLUE = "EAF2FB"
LIGHT_GRAY = "F6F7F9"
BORDER = "D9E2EC"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width: int, fnt, fill, line_gap: int = 8) -> int:
    x, y = xy
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline=None, radius=14, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_backend_home_mock(path: Path) -> None:
    img = Image.new("RGB", (1500, 1040), "#f4f6f8")
    d = ImageDraw.Draw(img)
    title_font = font(32, True)
    h_font = font(24, True)
    text_font = font(20)
    small_font = font(17)

    d.rectangle((0, 0, 1500, 72), fill="#101923")
    d.text((32, 20), "Sanity 后台 - 首页", font=title_font, fill="#ffffff")
    d.text((1290, 23), "Publish", font=font(22, True), fill="#ffffff")
    rounded(d, (1270, 14, 1428, 58), "#1f7a4d", radius=10)
    d.text((1304, 23), "Publish", font=font(22, True), fill="#ffffff")

    d.rectangle((0, 72, 300, 1040), fill="#ffffff")
    menu = ["站点设置", "固定页面", "  首页", "  新闻资讯页面", "  关于天赐宝贝", "新闻资讯", "科普视频中心", "咨询记录", "常见问题"]
    y = 110
    for item in menu:
        active = item.strip() == "首页"
        if active:
            rounded(d, (22, y - 8, 278, y + 38), "#e7f0ff", radius=8)
        d.text((42, y), item, font=text_font, fill="#155fa8" if active else "#344054")
        y += 54

    d.text((340, 112), "首页", font=font(36, True), fill="#111827")
    d.text((340, 162), "客户只需要重点看这几个模块：首屏、服务入口、优势、流程、新闻、FAQ、咨询区。", font=text_font, fill="#667085")

    modules = [
        ("专业首屏", "改首页第一屏大图、标题、说明、按钮、数据标签。", "最常改"),
        ("专业能力多图区", "改机构介绍、团队/实验室/咨询图片和能力说明。", "建议改"),
        ("快速服务入口", "改首页服务卡片的图片、标题、适用人群和链接。", "建议改"),
        ("核心优势", "改优势标题、说明、配图和排序。", "建议改"),
        ("助孕流程摘要", "改流程步骤、主图、按钮文案。", "建议改"),
        ("新闻 / FAQ / 咨询行动区", "控制首页底部内容、推荐文章和转化按钮。", "按需改"),
    ]
    y = 220
    for idx, (name, desc, tag) in enumerate(modules, start=1):
        rounded(d, (340, y, 1410, y + 88), "#ffffff", outline="#d0d7e2", radius=12, width=2)
        rounded(d, (365, y + 20, 420, y + 68), "#2465a8", radius=8)
        d.text((382, y + 32), str(idx), font=font(22, True), fill="#ffffff")
        d.text((445, y + 16), name, font=h_font, fill="#111827")
        d.text((445, y + 50), desc, font=small_font, fill="#667085")
        rounded(d, (1230, y + 24, 1355, y + 64), "#f3f8ff", outline="#bed4ef", radius=20)
        d.text((1254, y + 33), tag, font=small_font, fill="#2465a8")
        y += 106

    rounded(d, (340, 890, 1410, 980), "#fff7ed", outline="#fdba74", radius=12, width=2)
    d.text((370, 920), "客户动作：点左侧“固定页面 > 首页” -> 展开对应模块 -> 修改文字/图片 -> 点右上角 Publish。", font=font(22, True), fill="#a33d00")
    img.save(path)


def make_front_home_mock(path: Path) -> None:
    hero = ROOT / "public" / "images" / "home-hero.png"
    if hero.exists():
        base = Image.open(hero).convert("RGB").resize((1500, 940))
        overlay = Image.new("RGBA", base.size, (13, 23, 34, 115))
        img = Image.alpha_composite(base.convert("RGBA"), overlay).convert("RGB")
    else:
        img = Image.new("RGB", (1500, 940), "#d9e8ef")
    d = ImageDraw.Draw(img)

    rounded(d, (70, 52, 1430, 122), "#ffffff", radius=12)
    d.text((110, 73), "天赐宝贝", font=font(30, True), fill="#1f2937")
    for x, text in [(400, "首页"), (500, "试管服务"), (640, "医疗服务"), (790, "新闻资讯"), (930, "关于我们")]:
        d.text((x, 78), text, font=font(22), fill="#344054")
    rounded(d, (1230, 68, 1370, 108), "#2465a8", radius=20)
    d.text((1255, 78), "立即咨询", font=font(20, True), fill="#ffffff")

    d.text((105, 240), "专注辅助生殖服务", font=font(24, True), fill="#d8ecff")
    d.text((105, 295), "让每个家庭都能安心迎接新生命", font=font(48, True), fill="#ffffff")
    d.text((105, 370), "首页第一屏由 Sanity 后台“专业首屏”控制：大图、标题、说明、按钮、标签和数据都在这里维护。", font=font(25), fill="#eff8ff")
    rounded(d, (105, 450, 255, 506), "#ffffff", radius=28)
    d.text((132, 464), "立即咨询", font=font(22, True), fill="#1d4f84")
    rounded(d, (280, 450, 430, 506), "#2465a8", radius=28)
    d.text((307, 464), "了解更多", font=font(22, True), fill="#ffffff")

    labels = [("95%", "成功率"), ("20年", "行业经验"), ("3000+", "服务家庭")]
    x = 105
    for value, label in labels:
        rounded(d, (x, 620, x + 250, 720), "#ffffff", radius=18)
        d.text((x + 32, 640), value, font=font(34, True), fill="#2465a8")
        d.text((x + 32, 682), label, font=font(20), fill="#667085")
        x += 280

    callouts = [
        ((82, 210, 680, 535), "A", "专业首屏：标题、说明、按钮"),
        ((92, 600, 895, 744), "B", "数据标签：数值和说明"),
        ((1218, 62, 1380, 112), "C", "顶部咨询按钮"),
    ]
    for box, letter, text in callouts:
        d.rounded_rectangle(box, radius=16, outline="#ffb020", width=5)
        rounded(d, (box[0], box[1] - 48, box[0] + 300, box[1] - 8), "#ffb020", radius=16)
        d.text((box[0] + 14, box[1] - 40), f"{letter}. {text}", font=font(18, True), fill="#1f2937")
    img.save(path)


def make_module_front_mock(path: Path) -> None:
    img = Image.new("RGB", (1500, 940), "#fbfaf8")
    d = ImageDraw.Draw(img)
    d.text((80, 70), "首页下方常改模块位置示意", font=font(42, True), fill="#1f2937")
    d.text((80, 126), "这些模块不用每天改，但客户最容易问“后台哪个字段对应网站哪里”。", font=font(24), fill="#667085")

    cards = [
        ("专业能力多图区", "机构介绍、团队图、实验室图、咨询图", "#e8f4ff"),
        ("快速服务入口", "服务卡片图片、标题、适用人群、按钮链接", "#fff3e8"),
        ("核心优势", "优势列表、排序、配图、颜色", "#edf8ee"),
        ("助孕流程摘要", "流程步骤、主图、步骤编号", "#f2ecff"),
        ("新闻推荐", "首页展示哪些新闻，数量多少", "#f8eef2"),
        ("FAQ / 咨询行动区", "常见问题数量、底部咨询按钮", "#eef6f7"),
    ]
    positions = [(80, 210), (535, 210), (990, 210), (80, 535), (535, 535), (990, 535)]
    for idx, ((title, desc, fill), (x, y)) in enumerate(zip(cards, positions), start=1):
        rounded(d, (x, y, x + 390, y + 250), fill, outline="#d0d7e2", radius=18, width=2)
        rounded(d, (x + 24, y + 24, x + 92, y + 92), "#2465a8", radius=14)
        d.text((x + 47, y + 42), str(idx), font=font(28, True), fill="#ffffff")
        d.text((x + 24, y + 120), title, font=font(27, True), fill="#111827")
        draw_wrapped(d, desc, (x + 24, y + 164), 320, font(21), "#667085")
        d.rounded_rectangle((x + 18, y + 18, x + 372, y + 232), radius=18, outline="#ffb020", width=4)
    img.save(path)


def set_run(run, size=None, bold=None, color=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = node.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            node.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER)


def margins(cell, top=90, bottom=90, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = tc_mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def add_p(doc, text, size=10.5, bold=False, color=DARK, style=None, after=5):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table)
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run(r, size=10.2, bold=True, color=BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    headers = ["后台看到的字段", "网站对应位置", "客户怎么填"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, LIGHT_GRAY)
        margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = cell.paragraphs[0].add_run(header)
        set_run(r, size=9.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run(r, size=9.1, bold=(idx == 0), color=DARK if idx != 2 else MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    for style_name, size, color in [("Heading 1", 18, BLUE), ("Heading 2", 13, BLUE)]:
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    backend = OUT_DIR / "home-backend-mock.png"
    front = OUT_DIR / "home-front-hero-mock.png"
    modules = OUT_DIR / "home-front-modules-mock.png"
    make_backend_home_mock(backend)
    make_front_home_mock(front)
    make_module_front_mock(modules)

    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Sanity 后台首页截图对照样稿")
    set_run(r, size=24, bold=True, color=BLUE)
    add_p(doc, "给客户看的版本：只讲“在哪里改、网站哪里变、怎么填”，不讲后台技术名词。", size=11, color=MUTED, after=8)
    add_note(doc, "使用路径：Sanity 后台 -> 左侧“固定页面” -> “首页” -> 展开对应模块 -> 修改后点击右上角 Publish。")
    add_matrix(
        doc,
        [
            ["专业首屏", "首页打开后第一屏的大图、标题、按钮和数据标签", "最常改。图片先换，标题控制在 12-20 个字，按钮链接用 / 开头的站内地址。"],
            ["快速服务入口", "首页中部服务卡片", "一张卡对应一个服务。改标题、适用人群、说明、图片和按钮链接。"],
            ["核心优势", "展示机构优势的模块", "建议 3-5 条，能拖动排序。不要写太长，每条像一句卖点。"],
            ["新闻 / FAQ", "首页底部推荐内容", "新闻来自新闻管理，FAQ 来自常见问题。这里主要控制显示数量或手动推荐。"],
            ["咨询行动区", "首页底部的大咨询入口", "保留醒目的咨询按钮，电话/微信和按钮文案要检查。"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("一、后台首页在哪里", level=1)
    add_p(doc, "客户先看左侧菜单，不需要理解 Sanity 的文档类型；只要记住“固定页面 > 首页”。", size=10.5, color=MUTED)
    doc.add_picture(str(backend), width=Inches(6.9))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_matrix(
        doc,
        [
            ["左侧“固定页面 > 首页”", "进入首页所有可编辑内容", "客户每次改首页都从这里进。"],
            ["右上角 Publish", "让网站显示最新内容", "改完必须点；没点就是草稿，前台不会更新。"],
            ["模块标题", "对应首页不同区域", "找不到字段时，先找模块名字，不要逐个字段猜。"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("二、首页首屏 Hero 对照", level=1)
    add_p(doc, "这是客户最常改的区域。截图中的 A/B/C 对应后台“专业首屏”和站点设置里的按钮信息。", size=10.5, color=MUTED)
    doc.add_picture(str(front), width=Inches(6.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_matrix(
        doc,
        [
            ["专业首屏 > 桌面端大图 / 移动端大图", "A 区整张首屏背景图", "桌面建议 2400x1200，手机建议 900x1200，人物主体不要压在文字下面。"],
            ["专业首屏 > 主标题 / 专业说明", "A 区标题和说明文字", "标题短一点，说明写机构优势，不要堆太多专业术语。"],
            ["专业首屏 > 主按钮 / 次按钮", "A 区两个行动按钮", "按钮文字 4-6 个字，链接确认能打开。"],
            ["专业首屏 > 首屏数据栏", "B 区 95%、20年、3000+ 等数据", "每项包含“数值”和“说明”，最多 4 项。"],
            ["站点设置 / 顶部咨询信息", "C 区顶部咨询按钮", "确认电话、微信、按钮文字和跳转地址。"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("三、首页下方模块对照", level=1)
    add_p(doc, "下面这些模块不是每天改，但客户问“这个块在哪里改”时，可以直接对照编号。", size=10.5, color=MUTED)
    doc.add_picture(str(modules), width=Inches(6.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_matrix(
        doc,
        [
            ["1 专业能力多图区", "机构介绍和多张展示图", "适合放团队、实验室、咨询场景。图片和文字要匹配。"],
            ["2 快速服务入口", "服务卡片区域", "每张卡可拖动排序；标题、说明、链接都要检查。"],
            ["3 核心优势", "优势列表区域", "建议写真实优势，避免空泛词。"],
            ["4 助孕流程摘要", "流程步骤区域", "按步骤编号填写，保持 3-6 步最清楚。"],
            ["5 新闻推荐", "首页新闻区域", "需要先发布新闻文章，再在首页选择推荐。"],
            ["6 FAQ / 咨询行动区", "常见问题和底部咨询入口", "FAQ 先在常见问题里维护，咨询区负责引导成交。"],
        ],
    )

    add_note(doc, "这是一版样稿：后台图为结构示意，前台图为首页视觉示意。确认版式后，可用真实 Sanity 登录截图替换示意图。")
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
